import os
import pandas as pd
import yfinance as yf 
from google import genai
from google.genai import types
from docx import Document # 用於生成 Word 文件
from linebot import LineBotApi
from linebot.models import TextSendMessage

# --- 1. 配置 API Key 和 LINE 設定 ---
# 請將以下 YOUR_... 替換為你的實際金鑰和 ID
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "YOUR_GEMINI_API_KEY") 
LINE_CHANNEL_ACCESS_TOKEN = os.getenv("LINE_ACCESS_TOKEN", "YOUR_LINE_ACCESS_TOKEN")
LINE_USER_ID = "YOUR_LINE_USER_ID" # 接收訊息的 Line 用戶 ID 或群組 ID

# 初始化 Gemini 客戶端
try:
    client = genai.Client(api_key=GEMINI_API_KEY)
except Exception as e:
    print(f"Gemini 客戶端初始化失敗: {e}")
    client = None

# 初始化 Line 客戶端
try:
    line_bot_api = LineBotApi(LINE_CHANNEL_ACCESS_TOKEN)
except Exception as e:
    print(f"Line Bot 客戶端初始化失敗: {e}")
    line_bot_api = None


# --- 2. 數據獲取與分析準備 (與前一個版本相似) ---
def get_stock_analysis_data(ticker="2330.TW"):
    """
    使用 yfinance 獲取台積電的近期數據並進行簡單分析。
    """
    try:
        stock = yf.Ticker(ticker)
        history = stock.history(period="30d") 
        
        current_price = history['Close'].iloc[-1]
        previous_close = history['Close'].iloc[-2]
        change_percent = (current_price - previous_close) / previous_close * 100
        
        # 準備提供給 Gemini 分析的數據摘要
        analysis_summary = f"""
        股票代碼: {ticker} (台積電)
        近 30 日收盤價平均: {history['Close'].mean():.2f}
        今日收盤價: {current_price:.2f}
        與前一日比較漲跌幅: {change_percent:.2f}%
        
        --- [以下為範例，請替換成真實計算的技術指標結果] ---
        技術指標狀態 (需自行計算 RSI, MACD 等)：
        - 10日移動平均線 (MA10): 多頭排列 (假設)
        - 相對強弱指標 (RSI): 65 (中性偏強)
        - 買賣動能 (MACD): 紅柱增加 (動能轉強)
        
        市場情緒總結: 近期受 AI 晶片利多消息影響，整體情緒樂觀，成交量溫和放大。
        """
        return analysis_summary
    except Exception as e:
        return f"無法獲取股票數據或進行分析：{e}"


# --- 3. 使用 Gemini 整理和分析數據 ---
def analyze_with_gemini(data_summary):
    """
    將股票數據摘要提供給 Gemini 進行整理和趨勢預測。
    """
    if not client:
        return "Gemini API 未初始化，無法進行分析。"

    prompt = f"""
    請你扮演一位資深金融分析師。
    請根據以下台積電 (2330) 的近期數據和技術指標分析，
    為投資人整理出一份專業的報告。
    
    1. 報告應包含：**近期走勢總結**、**各項技術指標解讀**、**未來趨勢預測 (僅供參考)**。
    2. 請將結果整理成**報告格式**，包含標題和段落，以便存入 Word 文件。
    
    分析數據如下：
    {data_summary}
    """
    
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
        )
        return response.text
    except Exception as e:
        return f"Gemini API 呼叫失敗：{e}"

# --- 4. 儲存 Word 檔案 (.docx) ---
def save_to_word(content, filename="TSMC_Analysis_Report.docx"):
    """
    將 Gemini 產生的內容儲存為 Word 文件。
    """
    try:
        document = Document()
        document.add_heading('台積電 (2330) 專業分析報告', 0)
        
        # 將 Gemini 產生的多行文字分割並加入文件
        for line in content.split('\n'):
            line = line.strip()
            if line:
                # 簡單判斷，將粗體行視為次標題
                if line.startswith('**') and line.endswith('**'):
                    document.add_heading(line.replace('**', ''), level=1)
                else:
                    document.add_paragraph(line)
                    
        document.save(filename)
        return f"報告已成功儲存為 {filename}"
    except Exception as e:
        return f"儲存 Word 文件失敗: {e}"

# --- 5. 透過 LINE Messaging API 傳送檔案連結 ---
def send_line_file_link(file_name, public_url):
    """
    發送 Line 訊息，告知用戶檔案已生成並提供下載連結。
    """
    if not line_bot_api:
        return "Line Bot API 未初始化，無法發送訊息。"

    # 構造給 Line 用戶的訊息
    message = f"""
    【📢 檔案通知】台積電 (2330) 分析報告已生成！
    
    文件名稱: {file_name}
    
    請點擊下方連結下載 Word 報告 (.docx):
    {public_url}
    
    (請注意: 由於 LINE 限制，文件需上傳至雲端空間才能提供連結下載)
    """
    
    try:
        line_bot_api.push_message(
            LINE_USER_ID,
            TextSendMessage(text=message.strip())
        )
        return "Line 檔案連結訊息發送成功。"
    except Exception as e:
        return f"Line 訊息發送失敗：{e}"


# --- 主程式執行區塊 ---
def main():
    REPORT_FILENAME = "TSMC_Analysis_Report.docx"
    
    # --- 步驟 1: 獲取數據與分析 ---
    print("--- 步驟 1: 獲取台積電 (2330) 數據並使用 Gemini 分析 ---")
    stock_data = get_stock_analysis_data()
    gemini_report_content = analyze_with_gemini(stock_data)
    #print("Gemini 產生的報告內容:\n", gemini_report_content[:500] + "...") # 顯示部分內容
    
    # --- 步驟 2: 儲存 WORD 檔案 ---
    print("\n--- 步驟 2: 儲存 Word 檔案 (.docx) ---")
    word_save_result = save_to_word(gemini_report_content, REPORT_FILENAME)
    print(word_save_result)
    
    # --- 步驟 3: 透過 LINE 發送檔案連結 (重要提示) ---
    #print("\n--- 步驟 3: 透過 LINE 發送檔案連結 ---")
    
    # *** 這裡你需要手動操作或加入程式碼來將 REPORT_FILENAME 上傳到網路空間 ***
    # *** 假設你上傳後獲得了以下公開下載 URL ***
    
    # !!! 請將此處的範例連結替換為你實際的檔案公開下載連結 !!!
    PUBLIC_DOWNLOAD_URL = "https://your-cloud-storage.com/path/to/TSMC_Analysis_Report.docx" 
    
    if PUBLIC_DOWNLOAD_URL == "https://your-cloud-storage.com/path/to/TSMC_Analysis_Report.docx":
        print("⚠️ 警告：請將程式中的 PUBLIC_DOWNLOAD_URL 替換為實際的檔案公開連結。")
        print("    Line 訊息將使用範例連結發送，用戶無法成功下載。")

    line_send_result = send_line_file_link(REPORT_FILENAME, PUBLIC_DOWNLOAD_URL)
    print(line_send_result)

if __name__ == "__main__":
    main()