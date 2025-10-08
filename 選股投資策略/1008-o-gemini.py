import os
from google import genai
from google.genai.errors import APIError
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 設定 ---

# 檔案名稱 (Word 格式)
OUTPUT_FILENAME = "motech_news_summary.docx"

# 您的查詢訊息
PROMPT_MESSAGE = (
    "請幫我蒐集https://www.cnyes.com/的網站資訊，針對台灣股票茂矽2342的近期一星期的新聞，"
    "並整理出200字重點訊息提供。"
)

# --- 核心函式 ---

def save_to_word_file(query, content, filename):
    """
    將查詢訊息和 Gemini 的回覆內容儲存到 Word 文件中。
    """
    try:
        document = Document()
        
        # 標題樣式
        document.add_heading('Gemini 股票新聞摘要報告', 0) # level 0 是最大的標題

        # --- 寫入查詢訊息 ---
        document.add_heading('1. 查詢提示 (Prompt)', level=1)
        
        p_query = document.add_paragraph()
        # 設置字體大小和粗體
        run_query = p_query.add_run(query)
        run_query.bold = True
        font_query = run_query.font
        font_query.size = Pt(12)
        
        document.add_paragraph() # 空一行

        # --- 寫入 Gemini 回覆 ---
        document.add_heading('2. Gemini 摘要回覆 (茂矽 2342)', level=1)
        
        p_response = document.add_paragraph()
        run_response = p_response.add_run(content)
        # 設置標準字體
        font_response = run_response.font
        font_response.size = Pt(11)
        
        # 儲存文件
        document.save(filename)
        print(f"\n--- 完成 ---")
        print(f"回覆內容已成功儲存至 Word 檔案：{filename}")
        
    except Exception as e:
        print(f"\n--- 儲存 Word 檔案失敗 ---")
        print(f"無法寫入檔案 {filename}。錯誤：{e}")


def get_gemini_response_and_save():
    """
    連接 Gemini API，獲取回覆，並將結果傳遞給儲存函式。
    """
    try:
        # 初始化 Gemini 客戶端
        client = genai.Client()
    except Exception as e:
        print(f"初始化 Gemini 客戶端失敗：{e}")
        print("請確認您已設定 GEMINI_API_KEY 環境變數或在程式中指定金鑰。")
        return

    print("--- 正在呼叫 Gemini API，請稍候... ---")
    
    try:
        # 呼叫模型並啟用 Google 搜尋工具 (tool_config={'tools': [{'google_search': {}}]})
        # 這樣模型才能獲取最新的新聞資訊。
        response = client.models.generate_content(
            model='gemini-2.5-flash', 
            contents=PROMPT_MESSAGE,
            config={'tools': [{'google_search': {}}]} # 啟用 Google 搜尋以獲取最新資訊
        )

        # 檢查回覆內容
        if not response.text:
            print("Gemini 未提供回覆內容。")
            content_to_save = "Gemini 未提供回覆內容，請檢查 API 狀態或提示內容。"
        else:
            print("Gemini 回覆成功。")
            content_to_save = response.text

    except APIError as e:
        print(f"\n--- API 呼叫失敗 ---")
        content_to_save = f"API 呼叫失敗，錯誤：{e}"
    except Exception as e:
        print(f"\n--- 發生其他錯誤 ---")
        content_to_save = f"發生未知錯誤：{e}"

    # 將回覆內容儲存為 Word 檔案
    save_to_word_file(PROMPT_MESSAGE, content_to_save, OUTPUT_FILENAME)


if __name__ == "__main__":
    get_gemini_response_and_save()